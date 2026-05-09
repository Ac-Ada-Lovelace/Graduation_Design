from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "thesis" / "华北电力大学本科毕业设计论文_大纲版V2.docx"


def set_run_font(run, east_asia="宋体", latin="Times New Roman", size=None, bold=None, italic=None):
    run.font.name = latin
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)


def set_style_font(style, east_asia="宋体", latin="Times New Roman", size=12, bold=False):
    style.font.name = latin
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)


def set_section_layout(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.2)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = instruction
    run._r.append(instr)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)
    return run


def set_page_numbering(section, start=1, fmt="decimal"):
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))
    pg_num_type.set(qn("w:fmt"), fmt)


def set_center_footer_page(section, instruction):
    section.footer.is_linked_to_previous = False
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    add_field(paragraph, instruction)
    for run in paragraph.runs:
        set_run_font(run, size=9)


def set_body_header(section):
    section.header.is_linked_to_previous = False
    paragraph = section.header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    run = paragraph.add_run("华北电力大学本科毕业设计（论文）")
    set_run_font(run, east_asia="宋体", size=9)


def add_centered(doc, text, size=12, east_asia="宋体", bold=False, space_after=0):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_after = Pt(space_after)
    run = paragraph.add_run(text)
    set_run_font(run, east_asia=east_asia, size=size, bold=bold)
    return paragraph


def add_body_paragraph(doc, text, first_line=True):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Pt(24)
    run = paragraph.add_run(text)
    set_run_font(run, east_asia="宋体", size=12)
    return paragraph


def add_heading(doc, text, level=1, page_break=False):
    if page_break:
        doc.add_page_break()
    paragraph = doc.add_heading(text, level=level)
    if level == 1:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(12)
        size = 18
    elif level == 2:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)
        size = 15
    elif level == 3:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)
        size = 14
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)
        size = 12
    paragraph.paragraph_format.first_line_indent = None
    for run in paragraph.runs:
        set_run_font(run, east_asia="黑体", latin="Times New Roman", size=size, bold=False)
    return paragraph


def add_reference(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, east_asia="宋体", size=10.5)
    return paragraph


def add_caption(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, east_asia="宋体", size=10.5)
    return paragraph


def enable_update_fields_on_open(document):
    settings = document.settings.element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    else:
        update = existing
    update.set(qn("w:val"), "true")


def build():
    doc = Document()
    for section in doc.sections:
        set_section_layout(section)

    normal = doc.styles["Normal"]
    set_style_font(normal, east_asia="宋体", size=12)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.first_line_indent = Pt(24)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    for name, size in [("Heading 1", 18), ("Heading 2", 15), ("Heading 3", 14), ("Heading 4", 12)]:
        style = doc.styles[name]
        set_style_font(style, east_asia="黑体", size=size)
        style.paragraph_format.first_line_indent = None

    cover = doc.sections[0]
    set_section_layout(cover)
    cover.header.is_linked_to_previous = False
    cover.footer.is_linked_to_previous = False

    add_centered(doc, "毕 业 设 计（论 文）", size=22, east_asia="黑体", space_after=42)
    add_centered(doc, "论文题目：面向家庭能耗场景的 NILM 在线监测与推理系统设计与实现", size=14, east_asia="宋体", space_after=24)
    for label in ["院  系：", "专  业：", "班  级：", "姓  名：", "学  号：", "指导教师："]:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Pt(96)
        paragraph.paragraph_format.line_spacing = 1.5
        run = paragraph.add_run(label + " " * 18)
        set_run_font(run, east_asia="宋体", size=14)
    add_centered(doc, "年    月", size=14, east_asia="宋体", space_after=0)

    front = doc.add_section(WD_SECTION.NEW_PAGE)
    set_section_layout(front)
    set_page_numbering(front, start=1, fmt="upperRoman")
    set_center_footer_page(front, "PAGE \\* ROMAN")

    add_centered(doc, "摘    要", size=15, east_asia="黑体", bold=False, space_after=12)
    add_body_paragraph(doc, "此处撰写中文摘要。建议按“研究背景—问题定义—系统方法—实现内容—测试结果—应用价值”的顺序组织，突出本文是面向家庭能耗场景的 NILM 在线监测与推理系统，而不是单纯算法模型或单纯后台管理系统。摘要约400字，具有独立性和完整性。")
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run("关键词：")
    set_run_font(run, east_asia="宋体", size=12, bold=True)
    run = paragraph.add_run("非侵入式负荷监测；家庭能耗；数据采集；在线推理；系统设计与实现")
    set_run_font(run, east_asia="宋体", size=12)

    doc.add_page_break()
    add_centered(doc, "ABSTRACT", size=15, east_asia="黑体", bold=False, space_after=12)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(24)
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run("This section provides the English abstract corresponding to the Chinese abstract. It should describe the background, problem, system method, implementation modules, test results, and practical value of the NILM online monitoring and inference system.")
    set_run_font(run, east_asia="宋体", latin="Times New Roman", size=12)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run("KEY WORDS: ")
    set_run_font(run, latin="Times New Roman", size=12, bold=True)
    run = paragraph.add_run("non-intrusive load monitoring, household energy consumption, data acquisition, online inference, system design and implementation")
    set_run_font(run, latin="Times New Roman", size=12)

    doc.add_page_break()
    add_centered(doc, "目    录", size=15, east_asia="黑体", bold=False, space_after=12)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.line_spacing = 1.0
    add_field(paragraph, 'TOC \\o "1-3" \\h \\z \\u')
    add_body_paragraph(doc, "提示：首次打开本文档后，在 Word 中右键目录并选择“更新域/更新整个目录”。", first_line=False)

    body = doc.add_section(WD_SECTION.NEW_PAGE)
    set_section_layout(body)
    set_page_numbering(body, start=1, fmt="decimal")
    set_body_header(body)
    set_center_footer_page(body, "PAGE")

    chapters = [
        (
            "第1章 绪论",
            "本章解决“为什么做、别人做到了哪里、本文到底做什么”的问题。避免泛泛谈智能家居，重点落到 NILM 工程系统闭环不足。",
            [
                ("1.1 研究背景与意义", "说明家庭能耗管理、智能家居、节能减排和负荷识别需求，引出 NILM 的低侵入优势与工程落地价值。"),
                ("1.2 国内外研究现状", "按 NILM 方法、在线监测系统、工程部署三个方向综述，最后指出采集、推理、展示、验证衔接不足。"),
                ("1.3 研究内容与技术路线", "明确本文完成数据采集、协议解析、模型工件化、在线推理、应用展示和系统测试。"),
                ("1.4 本文主要工作", "用 3-4 点概括工程贡献，避免写成空泛创新点。"),
                ("1.5 论文组织结构", "简述各章安排。"),
            ],
        ),
        (
            "第2章 相关技术与需求分析",
            "本章为后续设计提供依据。相关技术不要写成百科，应只写本文会实际用到的技术。",
            [
                ("2.1 非侵入式负荷监测技术概述", "介绍 NILM 基本概念、输入输出、负荷分解、事件检测与评价指标。"),
                ("2.2 系统关键技术", "说明 TCP 通信、时间序列窗口、推理服务 API、前端可视化和在线回放等技术。"),
                ("2.3 系统功能需求分析", "从终端接入、报文解析、数据存储、离线推理、在线推理、结果展示六方面展开。"),
                ("2.4 系统非功能需求分析", "围绕稳定性、可维护性、可扩展性、可追溯性、可测试性写，后面要能被测试章节验证。"),
                ("2.5 本章小结", "收束为总体设计的输入。"),
            ],
        ),
        (
            "第3章 系统总体设计",
            "本章是全文的骨架，要把系统边界、模块划分和数据流讲清楚。这里可以保留三层思想，但要用实现证据支撑。",
            [
                ("3.1 系统设计目标", "明确目标是实现从数据采集到推理展示的闭环，而不是只完成模型训练。"),
                ("3.2 系统总体架构", "按数据采集层、NILM 推理服务层、应用展示层组织，并配总体架构图。"),
                ("3.3 数据流与业务流程设计", "分别说明采集数据流、离线推理流程、在线回放推理流程和前端展示流程。"),
                ("3.4 关键模块划分", "用表格列出数据采集、协议解析、数据存储、模型包管理、推理服务、在线回放、应用展示模块。"),
                ("3.5 数据接口与模型包契约设计", "写清输入字段、输出字段、模型包内容、版本信息和兼容性约束。"),
                ("3.6 本章小结", "说明第4-6章将分别展开各模块实现。"),
            ],
        ),
        (
            "第4章 数据采集与预处理模块设计实现",
            "本章对应 Stage-01，必须多写实现细节和运行效果，少写抽象原则。",
            [
                ("4.1 数据采集链路设计", "说明终端模拟器、TCP 监听、多终端连接管理和采集流程。"),
                ("4.2 通信协议与报文解析", "说明固定报文结构、字段含义、解析流程、单位换算和格式校验。"),
                ("4.3 设备状态管理与异常处理", "说明在线/离线判断、异常报文、断连、半包或格式错误处理。"),
                ("4.4 数据存储与追溯设计", "说明 JSONL 落盘、原始报文与解析结果保存，以及对后续训练和推理的数据支撑。"),
                ("4.5 采集模块运行效果", "放运行截图、设备列表截图、报文详情截图、采集记录样例。"),
                ("4.6 本章小结", "说明采集模块如何为推理服务提供可信输入。"),
            ],
        ),
        (
            "第5章 NILM 推理服务模块设计实现",
            "本章对应 Stage-02，是论文技术含量的核心之一，要写清模型如何变成可调用服务。",
            [
                ("5.1 数据处理与模型训练流程", "说明数据准备、窗口构造、训练/验证/测试划分和模型训练流程。"),
                ("5.2 模型工件与版本管理", "说明 model package、manifest、interface spec、registry、active package。"),
                ("5.3 离线推理服务设计", "说明区间推理、指标计算、预测与真值对齐、报告输出。"),
                ("5.4 在线会话推理机制", "说明 session start、ingest、latest、滚动窗口、buffer ready 和事件抽取。"),
                ("5.5 回放与集成验证机制", "说明 replay feeder、integration scripts、smoke check 和 runbook。"),
                ("5.6 本章小结", "说明推理服务如何把 NILM 模型转化为可调用、可验证、可回滚的工程能力。"),
            ],
        ),
        (
            "第6章 系统应用界面与功能实现",
            "本章容易写虚，要牢牢围绕已实现页面、接口和结果展示。不要扩写成完整 SaaS 平台愿景。",
            [
                ("6.1 应用层功能定位", "明确应用层消费推理服务结果，提供可视化和交互入口，不重新训练模型。"),
                ("6.2 后端接口封装", "对应 Stage-03 backend 或 Stage-02 service，说明 API 调用、数据聚合、状态转发。"),
                ("6.3 前端页面与可视化展示", "说明 Dashboard、Device Detail、System Status 等页面，并放页面截图。"),
                ("6.4 模型状态与事件结果展示", "说明模型版本、在线状态、功率曲线、事件流、设备识别结果等展示方式。"),
                ("6.5 本章小结", "说明应用界面如何完成系统闭环展示。"),
            ],
        ),
        (
            "第7章 系统测试与结果分析",
            "本章用于证明工作量和实际能力，尽量用表格、指标和截图说话。",
            [
                ("7.1 测试环境与数据说明", "列出硬件环境、软件环境、数据集、采样周期、测试脚本和模型版本。"),
                ("7.2 数据采集模块测试", "测试 TCP 连接、多终端接入、报文解析、异常处理和数据落盘完整性。"),
                ("7.3 推理服务功能测试", "测试模型包加载、离线推理接口、在线会话接口、错误处理和服务健康检查。"),
                ("7.4 模型识别与误差指标分析", "分析 MAE、RMSE、MeanDiff、MaxAbsDiff、Event Precision/Recall/F1。"),
                ("7.5 在线回放与端到端测试", "说明回放喂数、滚动推理、前端展示、状态轮询和事件展示是否形成闭环。"),
                ("7.6 测试结果讨论", "分析成功点、不足点和误差来源，包括数据质量、模型泛化、阈值设置、真实场景差异。"),
                ("7.7 本章小结", "归纳系统功能可用性和工程有效性。"),
            ],
        ),
    ]

    for idx, (chapter, intro, sections) in enumerate(chapters):
        add_heading(doc, chapter, level=1, page_break=idx > 0)
        add_body_paragraph(doc, intro)
        for section_title, note in sections:
            add_heading(doc, section_title, level=2)
            add_body_paragraph(doc, note)
        if idx == 0:
            add_caption(doc, "图1-1 技术路线示意图")
        if idx == 2:
            add_caption(doc, "图3-1 系统总体架构图")
            add_caption(doc, "表3-1 系统关键模块划分表")

    add_heading(doc, "结    论", level=1, page_break=True)
    add_body_paragraph(doc, "此处总结全文工作、系统实现效果、主要贡献、不足与后续改进方向。")

    add_heading(doc, "参 考 文 献", level=1, page_break=True)
    add_reference(doc, "[1] 作者. 文献题名[J]. 刊名, 年份, 卷号(期号): 起止页码.")
    add_reference(doc, "[2] 作者. 书名[M]. 出版地: 出版者, 出版年: 起止页码.")

    add_heading(doc, "附    录", level=1, page_break=True)
    add_heading(doc, "附录A 协议字段与偏移表", level=2)
    add_body_paragraph(doc, "此处放置协议字段、模型工件契约、API 样例、验收脚本与命令清单等补充材料。")

    add_heading(doc, "致    谢", level=1, page_break=True)
    add_body_paragraph(doc, "此处对导师、同学及在毕业设计过程中给予帮助的单位或个人表示感谢。")

    enable_update_fields_on_open(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
